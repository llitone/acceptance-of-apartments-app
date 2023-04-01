import os

from docx import Document
from docx.shared import Inches


class ReportGenerator(object):
    def __init__(self, filename: str, json: dict):
        self.filename = filename
        self.json = json
        if not os.path.isdir("db"):
            os.mkdir("db")
        if not os.path.isdir("db/reports"):
            os.mkdir("db/reports")
        self.doc = Document()
        header = self.doc.add_heading("ОТЧЕТ О НЕДОСТАТКАХ СТРОИТЕЛЬСТВА", 0)
        header.alignment = 1

    def save(self):
        self.doc.save("./db/reports/" + self.filename)

    def convert_json(self):
        i = 1
        for flat in self.json.items():
            for category in flat[1].items():
                category_header = self.doc.add_heading(f"{i}." + category[0], 1)
                category_header.bold = True
                for fault in category[1]:
                    text = self.doc.add_heading(
                        "\nПроблема: {0}\nОписание: {1}\nМесто: {2}\n".format(fault["name"], fault["description"],
                                                                              fault["place"]), 4)
                    count = 0
                    for image in fault["images"]:
                        picture = text.add_run("\n")
                        picture.add_picture(image, width=Inches(2))
                        count += 1
                    if count > 1:
                        self.doc.add_page_break()
                i += 1
# json_example = {
#     1337: {
#         "Личная жизнь":
#         [
#             {
#                 "name": "Московкин Александр",
#                 "place": "it-квадрат",
#                 "description": "МОСКВА ЖЕНАТ НА АЛЁНЕ И НЕ МОЖЕТ ЖЕНИТЬСЯ НА АНТОНЕ",
#                 "images":
#                     [
#                         "./db/images/jfuc8pKTFa4bCO0TbP5Tj3MkcEpvL7.jpg",
#                         "./db/images/MMKoFh78qxEMVz8M5iWA9NhfW3wl5T.jpg"
#                     ]
#             },
#             {
#                 "name": "Личная жизнь антона петлина",
#                 "place": "весь мир",
#                 "description": "её нет",
#                 "images": []
#             }
#         ],
#         "fuck": [
#             {
#                 "name": "Физика",
#                 "place": "школа",
#                 "description": "у максона 4...",
#                 "images": ["./db/images/MMKoFh78qxEMVz8M5iWA9NhfW3wl5T.jpg"]
#             }
#         ]
#     }
# }
#
# generator = ReportGenerator("output.docx", json_example)
# generator.convert_json()
# generator.save()
