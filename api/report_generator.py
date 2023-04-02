import os

from docx import Document
from docx.shared import Inches


class ReportGenerator(object):
    def __init__(self, filename: str, json: dict):
        self.filename = filename
        self.json = json
        if not os.path.isdir("db"):
            os.mkdir("db")
        if not os.path.isdir("db/reports"):
            os.mkdir("db/reports")
        self.doc = Document()
        header = self.doc.add_heading("ОТЧЕТ О НЕДОСТАТКАХ СТРОИТЕЛЬСТВА", 0)
        header.alignment = 1

    def save(self):
        self.doc.save("./db/reports/" + self.filename)

    def convert_json(self):
        i = 1
        for flat in self.json.items():
            for category in flat[1].items():
                category_header = self.doc.add_heading(f"{i}." + category[0], 1)
                category_header.bold = True
                for fault in category[1]:
                    text = self.doc.add_heading(
                        "\nПроблема: {0}\nМесто: {1}\n".format(fault["description"], fault["place"]), 4)
                    count = 0
                    for image in fault["images"]:
                        picture = text.add_run("\n")
                        picture.add_picture(image, width=Inches(2))
                        count += 1
                    if count > 1:
                        self.doc.add_page_break()
                i += 1
